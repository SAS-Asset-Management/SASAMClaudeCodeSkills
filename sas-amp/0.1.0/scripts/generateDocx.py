#!/usr/bin/env python3
"""
Generate a branded SAS-AM Asset Management Plan DOCX document.

Usage:
    python3 generateDocx.py \
        --sections-dir ./sas-amp-working/drafts/sections/ \
        --charts-dir ./sas-amp-working/data/charts/ \
        --output ./sas-amp-working/output/amp-document.docx \
        --org-name "Organisation Name" \
        --date "2026-03-23" \
        [--logo-path /path/to/logo.png]
"""

import argparse
import glob
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx not installed. Run: pip3 install python-docx")
    sys.exit(1)

# SAS-AM Brand Colours
SAS_BLUE = RGBColor(0x00, 0x22, 0x44)
SAS_GREEN = RGBColor(0x69, 0xBE, 0x28)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xD9, 0xD9, 0xD9)
DARK_TEXT = RGBColor(0x33, 0x33, 0x33)


def create_styles(doc):
    """Configure document styles with SAS-AM branding."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK_TEXT
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, size, spacing_before in [(1, 16, 24), (2, 14, 18), (3, 12, 12)]:
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(size)
        heading_style.font.bold = True
        heading_style.font.color.rgb = SAS_BLUE
        heading_style.paragraph_format.space_before = Pt(spacing_before)
        heading_style.paragraph_format.space_after = Pt(6)

    return doc


def add_cover_page(doc, org_name, date, logo_path=None):
    """Add branded cover page."""
    # Add spacing at top
    for _ in range(6):
        doc.add_paragraph()

    # Logo
    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(2))

    doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(org_name)
    run.font.size = Pt(28)
    run.font.color.rgb = SAS_BLUE
    run.font.bold = True

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Asset Management Plan')
    run.font.size = Pt(22)
    run.font.color.rgb = SAS_GREEN
    run.font.bold = False

    doc.add_paragraph()

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(date)
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_TEXT

    doc.add_paragraph()

    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run(
        'Prepared by SAS Asset Management\n'
        'We provide advanced analytics, expert asset management services '
        'and maturity assessments to help asset owners realise their value.'
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.font.italic = True

    # Page break
    doc.add_page_break()


def add_document_control(doc, org_name, date):
    """Add document control table."""
    doc.add_heading('Document Control', level=1)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    data = [
        ('Document Title', f'{org_name} — Asset Management Plan'),
        ('Date', date),
        ('Version', '1.0'),
        ('Status', 'Draft'),
        ('Prepared By', 'SAS Asset Management'),
    ]

    for i, (label, value) in enumerate(data):
        row = table.rows[i]
        # Label cell
        cell = row.cells[0]
        cell.text = label
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="002244"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE

        # Value cell
        row.cells[1].text = value
        for paragraph in row.cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(4)

    doc.add_paragraph()

    # Revision history table
    doc.add_heading('Revision History', level=2)
    rev_table = doc.add_table(rows=2, cols=5)
    rev_table.style = 'Table Grid'

    headers = ['Rev', 'Date', 'Prepared By', 'Reviewed By', 'Approved By']
    for i, header in enumerate(headers):
        cell = rev_table.rows[0].cells[i]
        cell.text = header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="002244"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)

    # First revision row
    rev_data = ['1.0', date, 'SAS Asset Management', '', '']
    for i, val in enumerate(rev_data):
        rev_table.rows[1].cells[i].text = val
        for paragraph in rev_table.rows[1].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

    doc.add_page_break()


def add_toc(doc):
    """Add a table of contents field."""
    doc.add_heading('Table of Contents', level=1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar)

    run = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run._r.append(instrText)

    run = paragraph.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run._r.append(fldChar)

    run = paragraph.add_run('(Right-click and select "Update Field" to populate)')
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True

    run = paragraph.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar)

    doc.add_page_break()


def parse_markdown_to_docx(doc, markdown_text, charts_dir=None):
    """Convert markdown section content to DOCX elements."""
    lines = markdown_text.split('\n')
    in_table = False
    table_rows = []
    in_list = False
    in_code_block = False

    for line in lines:
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            if in_table:
                _flush_table(doc, table_rows)
                table_rows = []
                in_table = False
            in_list = False
            continue

        # Code blocks
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph(stripped)
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            continue

        # Headings
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)

        # Table rows
        elif stripped.startswith('|') and stripped.endswith('|'):
            # Skip separator rows
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_rows.append(cells)
            in_table = True

        # Bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_run(p, text)

        # Numbered list
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_run(p, text)

        # Image reference
        elif stripped.startswith('!['):
            match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
            if match and charts_dir:
                alt_text, img_path = match.groups()
                full_path = os.path.join(charts_dir, os.path.basename(img_path))
                if os.path.exists(full_path):
                    doc.add_picture(full_path, width=Inches(5.5))
                    caption = doc.add_paragraph(alt_text)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style = doc.styles['Normal']
                    for run in caption.runs:
                        run.font.italic = True
                        run.font.size = Pt(9)

        # Regular paragraph
        else:
            p = doc.add_paragraph()
            _add_formatted_run(p, stripped)

    # Flush any remaining table
    if in_table and table_rows:
        _flush_table(doc, table_rows)


def _add_formatted_run(paragraph, text):
    """Add a run with basic markdown formatting (bold, italic)."""
    # Simple bold/italic parsing
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        else:
            paragraph.add_run(part)


def _flush_table(doc, rows):
    """Convert accumulated table rows into a DOCX table."""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

                # Style header row
                if i == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="002244"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = WHITE
                            run.font.bold = True

    doc.add_paragraph()


def add_back_page(doc):
    """Add back page with SAS-AM branding."""
    doc.add_page_break()

    for _ in range(8):
        doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('SAS Asset Management')
    run.font.size = Pt(20)
    run.font.color.rgb = SAS_BLUE
    run.font.bold = True

    doc.add_paragraph()

    mission = doc.add_paragraph()
    mission.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mission.add_run(
        'We provide advanced analytics, expert asset management services\n'
        'and maturity assessments to help asset owners realise their value.'
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.font.italic = True

    doc.add_paragraph()

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run('sas-am.com')
    run.font.size = Pt(11)
    run.font.color.rgb = SAS_GREEN


def main():
    parser = argparse.ArgumentParser(description='Generate SAS-AM branded AMP DOCX')
    parser.add_argument('--sections-dir', required=True, help='Directory containing section markdown files')
    parser.add_argument('--charts-dir', default=None, help='Directory containing chart images')
    parser.add_argument('--output', required=True, help='Output DOCX file path')
    parser.add_argument('--org-name', required=True, help='Organisation name')
    parser.add_argument('--date', required=True, help='Document date')
    parser.add_argument('--logo-path', default=None, help='Path to logo image')
    args = parser.parse_args()

    # Create document
    doc = Document()
    doc = create_styles(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Build document
    add_cover_page(doc, args.org_name, args.date, args.logo_path)
    add_document_control(doc, args.org_name, args.date)
    add_toc(doc)

    # Add sections in order
    section_files = sorted(glob.glob(os.path.join(args.sections_dir, '*.md')))

    if not section_files:
        print(f"WARNING: No section files found in {args.sections_dir}")

    for section_file in section_files:
        print(f"Processing: {os.path.basename(section_file)}")
        with open(section_file, 'r') as f:
            content = f.read()
        parse_markdown_to_docx(doc, content, args.charts_dir)

    # Add back page
    add_back_page(doc)

    # Ensure output directory exists
    os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)

    # Save
    doc.save(args.output)
    file_size = os.path.getsize(args.output)
    print(f"Document saved: {args.output} ({file_size:,} bytes)")


if __name__ == '__main__':
    main()
